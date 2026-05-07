from __future__ import annotations

import io
import math
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


CRM_DOC = Path(r"C:\Users\User\Documents\UNICEF\coisas\maza-crm\docs\MAZA_CRM_DOCUMENTATION_FINAL_COMPLETE.docx")
CRM_OUT = Path(r"C:\Users\User\Documents\UNICEF\coisas\maza-crm\docs\MAZA_CRM_DOCUMENTATION_APP_STYLE.docx")
BACKEND_DOC = Path(
    r"C:\Users\User\Documents\UNICEF\coisas\maza-strapi-backend\documentation\MAZA_STRAPI_BACKEND_DOCUMENTATION_FINAL_COMPLETE.docx"
)
BACKEND_OUT = Path(
    r"C:\Users\User\Documents\UNICEF\coisas\maza-strapi-backend\documentation\MAZA_STRAPI_BACKEND_DOCUMENTATION_APP_STYLE.docx"
)


def get_font(size: int, bold: bool = False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        proposal = word if not current else f"{current} {word}"
        if draw.textlength(proposal, font=font) <= width:
            current = proposal
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def new_canvas(title: str, subtitle: str, size=(1600, 900)):
    img = Image.new("RGB", size, (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, bold=True)
    subtitle_font = get_font(18)
    draw.text((60, 40), title, font=title_font, fill=(15, 23, 42))
    draw.text((60, 90), subtitle, font=subtitle_font, fill=(71, 85, 105))
    return img, draw


def draw_box(draw, xy, text, fill, outline, font, text_fill=(30, 41, 59), radius=18):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    lines = wrap_text(draw, text, font, max(80, x2 - x1 - 24))
    line_height = getattr(font, "size", 20) + 4
    total_h = len(lines) * line_height
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=font)
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=text_fill)
        y += line_height


def arrow(draw, start, end, fill=(42, 109, 186), width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 14
    spread = math.pi / 7
    p1 = (
        end[0] - head_len * math.cos(angle - spread),
        end[1] - head_len * math.sin(angle - spread),
    )
    p2 = (
        end[0] - head_len * math.cos(angle + spread),
        end[1] - head_len * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=fill)


def label(draw, pos, text, font, fill=(42, 109, 186)):
    w = draw.textlength(text, font=font)
    x, y = pos
    padding = 8
    h = getattr(font, "size", 18) + 10
    draw.rounded_rectangle((x, y, x + w + padding * 2, y + h), radius=12, fill=(235, 244, 255), outline=fill)
    draw.text((x + padding, y + 5), text, font=font, fill=fill)


def img_bytes(img: Image.Image) -> bytes:
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    return buffer.getvalue()


def make_crm_architecture() -> bytes:
    img, draw = new_canvas(
        "MAZA CRM Architecture",
        "CRM operators use a Vite frontend to manage content and analytics through the Strapi backend.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    draw_box(draw, (100, 160, 1490, 340), "CRM Users", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (150, 205, 430, 300), "CRM Admin", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (530, 205, 850, 300), "Content Manager", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (950, 205, 1380, 300), "Operations / Reporting User", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (320, 400, 1280, 580), "CRM Frontend", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (560, 450, 1040, 540), "Maza CRM\nReact + TypeScript + Vite", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (60, 660, 860, 860), "Backend Services", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (520, 710, 900, 800), "Maza Strapi Backend", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (100, 740, 330, 840), "Authentication /\nAuthorization", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (380, 740, 650, 840), "Content Management APIs", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (700, 740, 1000, 840), "Analytics & Reporting APIs", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (910, 660, 1540, 860), "Infrastructure & Integrations", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (970, 740, 1120, 830), "Database", (220, 252, 231), (20, 184, 166), small)
    draw_box(draw, (1160, 740, 1350, 830), "Media Storage", (220, 252, 231), (20, 184, 166), small)
    draw_box(draw, (1390, 740, 1510, 830), "PostHog /\nAnalytics", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (1120, 845, 1360, 915), "SMS / Email Services", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (1390, 845, 1560, 915), "External Integrations", (248, 250, 252), (203, 213, 225), small)
    arrow(draw, (290, 300), (630, 450))
    arrow(draw, (690, 300), (800, 450))
    arrow(draw, (1160, 300), (920, 450))
    arrow(draw, (800, 540), (760, 710))
    arrow(draw, (650, 760), (330, 790))
    arrow(draw, (720, 800), (600, 800))
    arrow(draw, (820, 780), (830, 800))
    arrow(draw, (900, 740), (1045, 775))
    arrow(draw, (900, 745), (1230, 775))
    arrow(draw, (900, 755), (1450, 775))
    arrow(draw, (860, 790), (1215, 880))
    return img_bytes(img)


def make_crm_login_flow() -> bytes:
    img, draw = new_canvas(
        "CRM Login Flow",
        "The CRM authenticates staff through the dedicated /auth/crm endpoint and stores a JWT for dashboard access.",
    )
    font = get_font(24, bold=True)
    boxes = [
        ((90, 350, 340, 490), "1\nUser opens\nlogin page"),
        ((400, 350, 650, 490), "2\nCredentials\nsubmitted"),
        ((710, 350, 1010, 490), "3\nPOST /auth/crm\nrequest"),
        ((1070, 350, 1320, 490), "4\nJWT + role\nreturned"),
        ((1380, 350, 1590, 490), "5\nToken stored and\ndashboard opens"),
    ]
    for coords, text in boxes:
        draw_box(draw, coords, text, (255, 255, 255), (59, 130, 246), font)
    for idx in range(len(boxes) - 1):
        arrow(draw, (boxes[idx][0][2], 420), (boxes[idx + 1][0][0], 420))
    return img_bytes(img)


def make_crm_course_flow() -> bytes:
    img, draw = new_canvas(
        "CRM Course Creation Flow",
        "Course authoring moves from initial setup through curriculum design, optional testing, draft review, and backend creation.",
    )
    font = get_font(24, bold=True)
    top = [
        ((90, 260, 340, 400), "1\nOpen course\ncreation"),
        ((420, 260, 670, 400), "2\nEnter basic\ninformation"),
        ((750, 260, 1000, 400), "3\nBuild modules\nand content"),
        ((1080, 260, 1330, 400), "4\nAdd final test\nif needed"),
    ]
    bottom = [
        ((1080, 600, 1330, 740), "5\nCourse ready for\npublish later"),
        ((750, 600, 1000, 740), "6\nUpload media via\nupload queue"),
        ((420, 600, 670, 740), "7\nCreate course\nin backend"),
        ((90, 600, 340, 740), "8\nReview\ndraft"),
    ]
    for coords, text in top + bottom:
        draw_box(draw, coords, text, (255, 255, 255), (59, 130, 246), font)
    for idx in range(len(top) - 1):
        arrow(draw, (top[idx][0][2], 330), (top[idx + 1][0][0], 330))
    arrow(draw, (1205, 400), (1205, 600))
    for idx in range(len(bottom) - 1):
        arrow(draw, (bottom[idx][0][0], 670), (bottom[idx + 1][0][2], 670))
    return img_bytes(img)


def make_crm_user_admin_flow() -> bytes:
    img, draw = new_canvas(
        "CRM User Administration Flow",
        "Administrators move from filtered user lists to account actions, then refresh the query-backed table.",
    )
    font = get_font(24, bold=True)
    boxes = [
        ((80, 350, 370, 490), "1\nOpen user\nmanagement"),
        ((450, 350, 740, 490), "2\nSearch and filter\nresult set"),
        ((820, 350, 1110, 490), "3\nCreate, edit, block,\nunblock, or delete"),
        ((1190, 350, 1520, 490), "4\nInvalidate query cache\nand refresh table"),
    ]
    for coords, text in boxes:
        draw_box(draw, coords, text, (255, 255, 255), (59, 130, 246), font)
    for idx in range(len(boxes) - 1):
        arrow(draw, (boxes[idx][0][2], 420), (boxes[idx + 1][0][0], 420))
    return img_bytes(img)


def make_crm_deployment() -> bytes:
    img, draw = new_canvas(
        "CRM Deployment Topology",
        "The CRM is served as a static web frontend and talks to the Strapi API configured through environment variables.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    draw_box(draw, (90, 360, 300, 470), "Browser Users", (255, 255, 255), (100, 116, 139), font)
    draw_box(draw, (390, 350, 690, 490), "Vercel / Static SPA\nMaza CRM Frontend", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (800, 350, 1100, 490), "Strapi Backend API", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (1190, 320, 1370, 420), "Database", (220, 252, 231), (20, 184, 166), font)
    draw_box(draw, (1190, 470, 1370, 570), "Media Storage", (220, 252, 231), (20, 184, 166), font)
    draw_box(draw, (1450, 320, 1590, 420), "Analytics /\nEmail / SMS", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (430, 640, 670, 740), "Local Dev:\npnpm dev + Vite", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (760, 640, 1000, 740), "Env Config:\nVITE_API_URL", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (1090, 640, 1380, 740), "Production:\nhttps://strapi.mazas.org/api", (255, 255, 255), (100, 116, 139), small)
    arrow(draw, (300, 415), (390, 415))
    arrow(draw, (690, 420), (800, 420))
    arrow(draw, (1100, 395), (1190, 370))
    arrow(draw, (1100, 430), (1190, 520))
    arrow(draw, (1100, 405), (1450, 370))
    arrow(draw, (540, 490), (560, 640))
    arrow(draw, (900, 490), (880, 640))
    arrow(draw, (1030, 490), (1235, 640))
    return img_bytes(img)


def make_backend_architecture() -> bytes:
    img, draw = new_canvas(
        "MAZA Strapi Backend Architecture",
        "The backend serves mobile, CRM, and admin clients while coordinating custom content, auth, services, and integrations.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    draw_box(draw, (230, 140, 1370, 290), "Platform Clients", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (300, 180, 520, 255), "Mobile App", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (600, 180, 810, 255), "CRM Frontend", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (900, 180, 1160, 255), "Strapi Admin UI", (255, 255, 255), (100, 116, 139), small)
    draw_box(draw, (80, 340, 1480, 620), "Maza Strapi Backend Core", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (640, 390, 960, 480), "Maza Strapi Backend", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (120, 470, 360, 560), "Users & Permissions\nExtension", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (430, 470, 650, 560), "Custom Content\nTypes", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (720, 470, 920, 560), "Lifecycle Hooks", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (990, 470, 1220, 560), "Custom Controllers\n& Routes", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (1290, 470, 1440, 560), "Services", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (340, 700, 1080, 860), "External Integrations", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (1180, 700, 1520, 860), "Infrastructure", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (430, 760, 580, 835), "SMS Providers", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (640, 760, 760, 835), "PostHog", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (820, 760, 940, 835), "Yoma API", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (980, 750, 1180, 845), "Google Play /\nApple App Store", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (410, 840, 620, 910), "OpenAI-Compatible API", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (1230, 760, 1430, 835), "Database\n(PostgreSQL / SQLite)", (220, 252, 231), (20, 184, 166), small)
    draw_box(draw, (1450, 760, 1560, 835), "Media\nStorage", (220, 252, 231), (20, 184, 166), small)
    arrow(draw, (410, 255), (700, 390))
    arrow(draw, (705, 255), (780, 390))
    arrow(draw, (1030, 255), (880, 390))
    for x in [240, 540, 820, 1105, 1365]:
        arrow(draw, (800, 480), (x, 470))
    arrow(draw, (1380, 560), (500, 760))
    arrow(draw, (1380, 560), (700, 760))
    arrow(draw, (1380, 560), (880, 760))
    arrow(draw, (1380, 560), (1080, 760))
    arrow(draw, (1380, 560), (500, 840))
    arrow(draw, (1380, 560), (1330, 760))
    arrow(draw, (1220, 520), (1500, 760))
    return img_bytes(img)


def make_backend_auth_flow() -> bytes:
    img, draw = new_canvas(
        "Backend Authentication and Access Flow",
        "Mobile and CRM clients authenticate through custom endpoints layered on top of the Strapi users-permissions extension.",
    )
    font = get_font(24, bold=True)
    draw_box(draw, (80, 330, 290, 430), "Mobile App", (255, 255, 255), (100, 116, 139), font)
    draw_box(draw, (80, 500, 290, 600), "CRM Frontend", (255, 255, 255), (100, 116, 139), font)
    draw_box(draw, (400, 320, 620, 430), "POST /auth/login", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (400, 500, 620, 610), "POST /auth/crm", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (730, 390, 1020, 540), "Users & Permissions\nExtension", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (1120, 240, 1360, 340), "Role / status\nvalidation", (219, 234, 254), (96, 165, 250), font)
    draw_box(draw, (1120, 400, 1360, 500), "JWT issuance", (219, 234, 254), (96, 165, 250), font)
    draw_box(draw, (1120, 560, 1360, 680), "GET /users/me and\nprotected routes", (219, 234, 254), (96, 165, 250), font)
    draw_box(draw, (1440, 380, 1570, 540), "Authenticated\naccess", (220, 252, 231), (20, 184, 166), font)
    arrow(draw, (290, 380), (400, 380))
    arrow(draw, (290, 550), (400, 550))
    arrow(draw, (620, 380), (730, 445))
    arrow(draw, (620, 550), (730, 485))
    arrow(draw, (1020, 450), (1120, 290))
    arrow(draw, (1020, 465), (1120, 450))
    arrow(draw, (1020, 485), (1120, 620))
    arrow(draw, (1360, 450), (1440, 460))
    arrow(draw, (1360, 620), (1440, 470))
    return img_bytes(img)


def make_backend_content_model() -> bytes:
    img, draw = new_canvas(
        "Backend Content Model Relationships",
        "Learning structure flows from journeys and categories into courses, modules, content, and learner progress records.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    top = [
        ((120, 280, 320, 370), "Journey"),
        ((420, 280, 620, 370), "Subject /\nCategory"),
        ((720, 280, 920, 370), "Course"),
        ((1020, 280, 1220, 370), "Course Module"),
        ((1320, 280, 1540, 370), "Content / Quiz"),
    ]
    for coords, text in top:
        draw_box(draw, coords, text, (219, 234, 254), (96, 165, 250), font)
    for idx in range(len(top) - 1):
        arrow(draw, (top[idx][0][2], 325), (top[idx + 1][0][0], 325))
    draw_box(draw, (720, 500, 960, 620), "User-Course\nProgress", (224, 231, 255), (79, 70, 229), font)
    arrow(draw, (820, 370), (820, 500))
    bottom = [
        ((380, 720, 560, 800), "Review"),
        ((700, 720, 900, 800), "Certificate"),
        ((1020, 720, 1210, 800), "Event / Analytics"),
        ((1320, 720, 1510, 800), "OTP / Identity"),
    ]
    for coords, text in bottom:
        draw_box(draw, coords, text, (248, 250, 252), (203, 213, 225), small)
    arrow(draw, (780, 620), (470, 720))
    arrow(draw, (820, 620), (800, 720))
    arrow(draw, (900, 620), (1115, 720))
    arrow(draw, (940, 580), (1415, 720))
    return img_bytes(img)


def make_backend_deployment() -> bytes:
    img, draw = new_canvas(
        "Backend Deployment Topology",
        "The observed production setup fronts the Strapi app with Nginx and connects it to storage, database, and messaging services.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    draw_box(draw, (80, 360, 280, 470), "Mobile / CRM /\nAdmin Clients", (255, 255, 255), (100, 116, 139), font)
    draw_box(draw, (380, 280, 1020, 720), "Hetzner Ubuntu Host", (241, 245, 249), (148, 163, 184), font)
    draw_box(draw, (450, 390, 700, 490), "Nginx / Public Host", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (730, 390, 990, 490), "Strapi App\n/var/www/maza-strapi-backend", (224, 231, 255), (79, 70, 229), font)
    draw_box(draw, (500, 560, 690, 660), "Strapi Admin UI", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (740, 560, 930, 660), "API Endpoints", (219, 234, 254), (96, 165, 250), small)
    draw_box(draw, (1140, 340, 1320, 440), "PostgreSQL", (220, 252, 231), (20, 184, 166), font)
    draw_box(draw, (1140, 510, 1320, 610), "Media / Uploads", (220, 252, 231), (20, 184, 166), font)
    draw_box(draw, (1140, 680, 1320, 780), "SMS / Email /\nAnalytics APIs", (248, 250, 252), (203, 213, 225), small)
    draw_box(draw, (1400, 470, 1560, 590), "Observed Production:\nhttps://strapi.maza...", (255, 255, 255), (100, 116, 139), small)
    arrow(draw, (280, 415), (450, 440))
    arrow(draw, (700, 440), (730, 440))
    arrow(draw, (990, 430), (1140, 390))
    arrow(draw, (990, 450), (1140, 560))
    arrow(draw, (990, 470), (1140, 730))
    arrow(draw, (830, 490), (790, 560))
    arrow(draw, (850, 490), (850, 560))
    arrow(draw, (1320, 390), (1400, 505))
    arrow(draw, (1320, 560), (1400, 525))
    return img_bytes(img)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="BFC7D5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(30, 64, 175)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor(51, 65, 85)

    if "Subtitle" not in styles:
        subtitle = styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(100, 116, 139)

    if "Caption" not in styles:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(71, 85, 105)

    for style_name in ["List Bullet", "List Number"]:
        if style_name in styles:
            styles[style_name].font.name = "Aptos"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
            styles[style_name].font.size = Pt(10.5)


def iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Unsupported parent")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='blip']"))


def get_paragraph_image_blob(paragraph: Paragraph):
    blips = paragraph._p.xpath(".//*[local-name()='blip']")
    if not blips:
        return None
    rel_id = blips[0].get(qn("r:embed"))
    if not rel_id:
        return None
    return paragraph.part.related_parts[rel_id].blob


def add_picture_from_blob(document: Document, blob: bytes, width=Inches(6.7)):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(blob), width=width)


def copy_paragraph(src: Paragraph, document: Document):
    text = src.text
    style_name = src.style.name if src.style else "Normal"
    is_caption = text.strip().startswith("Figure:")
    is_subtitle = style_name == "Subtitle" or text.strip().startswith("Generated from project documentation on")
    if not text.strip():
        document.add_paragraph("")
        return
    if style_name == "Title":
        p = document.add_paragraph(text, style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if is_subtitle:
        p = document.add_paragraph(text, style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if is_caption:
        p = document.add_paragraph(text, style="Caption")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if style_name.startswith("Heading 1") or style_name == "Heading1":
        document.add_paragraph(text, style="Heading 1")
        return
    if style_name.startswith("Heading 2") or style_name == "Heading2":
        document.add_paragraph(text, style="Heading 2")
        return
    if style_name.startswith("Heading 3") or style_name == "Heading3":
        document.add_paragraph(text, style="Heading 3")
        return
    if style_name.startswith("List Bullet"):
        document.add_paragraph(text, style="List Bullet")
        return
    if style_name.startswith("List Number"):
        document.add_paragraph(text, style="List Number")
        return
    p = document.add_paragraph(text)
    p.style = document.styles["Normal"]


def copy_table(src: Table, document: Document):
    rows = len(src.rows)
    cols = max(len(row.cells) for row in src.rows)
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(src.rows):
        for c_idx in range(cols):
            text = row.cells[c_idx].text if c_idx < len(row.cells) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=r_idx == 0)
            set_cell_border(table.cell(r_idx, c_idx))
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), "DCE6F1")


def format_doc(source_path: Path, output_path: Path, diagram_blobs: list[bytes]):
    src = Document(source_path)
    document = Document()
    configure_styles(document)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        add_page_number(section)

    for section in src.sections[1:]:
        new_section = document.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = section.orientation
        new_section.page_width = section.page_width
        new_section.page_height = section.page_height
        new_section.top_margin = Inches(0.7)
        new_section.bottom_margin = Inches(0.7)
        new_section.left_margin = Inches(0.8)
        new_section.right_margin = Inches(0.8)
        add_page_number(new_section)

    diagram_index = 0
    for block in iter_block_items(src):
        if isinstance(block, Paragraph):
            if paragraph_has_image(block):
                blob = get_paragraph_image_blob(block)
                if diagram_index < len(diagram_blobs):
                    blob = diagram_blobs[diagram_index]
                diagram_index += 1
                if blob is not None:
                    add_picture_from_blob(document, blob)
            else:
                copy_paragraph(block, document)
        else:
            copy_table(block, document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    crm_diagrams = [
        make_crm_architecture(),
        make_crm_login_flow(),
        make_crm_course_flow(),
        make_crm_user_admin_flow(),
        make_crm_deployment(),
    ]
    backend_diagrams = [
        make_backend_architecture(),
        make_backend_auth_flow(),
        make_backend_content_model(),
        make_backend_deployment(),
    ]
    format_doc(CRM_DOC, CRM_OUT, crm_diagrams)
    format_doc(BACKEND_DOC, BACKEND_OUT, backend_diagrams)
    print(CRM_OUT)
    print(BACKEND_OUT)


if __name__ == "__main__":
    main()
