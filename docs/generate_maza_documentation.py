from __future__ import annotations

import json
import math
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\User\Documents\UNICEF\coisas\org.maza.app")
BACKEND_ROOT = Path(r"C:\Users\User\Documents\UNICEF\maza-strapi-backend")
CRM_ROOT = Path(r"C:\Users\User\Documents\UNICEF\maza-crm")
SOURCE_DOC = Path(r"C:\Users\User\Documents\MAZA\Maza App Documentation V.01.docx")
OUT_DIR = ROOT / "docs" / "generated"
OUT_DOC = OUT_DIR / "Maza_App_Documentation_V04.docx"
DIAGRAM_DIR = OUT_DIR / "diagrams"

DISPLAY_ROOT = r"...\org.maza.app"
DISPLAY_BACKEND = r"...\maza-strapi-backend"
DISPLAY_CRM = r"...\maza-crm"
DISPLAY_SOURCE_DOC = r"...\Maza App Documentation V.01.docx"
DISPLAY_OUT_DOC = r"...\org.maza.app\docs\generated\Maza_App_Documentation_V04.docx"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)


def list_routes() -> list[str]:
    app_dir = ROOT / "app"
    routes = []
    for path in sorted(app_dir.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix not in {".tsx", ".ts"}:
            continue
        if any(part in {"components", "hooks", "services", "styles", "utils", "types", "__tests__"} for part in path.parts):
            continue
        rel = path.relative_to(app_dir).as_posix()
        routes.append(rel)
    return routes


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        ("arialbd.ttf" if bold else "arial.ttf"),
        ("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"),
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        proposal = word if not current else f"{current} {word}"
        if draw.textlength(proposal, font=font) <= width:
            current = proposal
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, fill, outline, font, text_fill=(30, 41, 59), radius=18):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    lines = wrap_text(draw, text, font, max(80, x2 - x1 - 24))
    total_h = len(lines) * (font.size + 4)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=font)
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=text_fill)
        y += font.size + 4


def arrow(draw, start, end, fill=(42, 109, 186), width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 14
    spread = math.pi / 7
    p1 = (
        end[0] - head_len * math.cos(angle - spread),
        end[1] - head_len * math.sin(angle - spread),
    )
    p2 = (
        end[0] - head_len * math.cos(angle + spread),
        end[1] - head_len * math.sin(angle + spread),
    )
    draw.polygon([end, p1, p2], fill=fill)


def label(draw, pos, text, font, fill=(42, 109, 186)):
    w = draw.textlength(text, font=font)
    x, y = pos
    padding = 8
    draw.rounded_rectangle((x, y, x + w + padding * 2, y + font.size + 10), radius=12, fill=(235, 244, 255), outline=fill)
    draw.text((x + padding, y + 5), text, font=font, fill=fill)


def new_canvas(title: str, subtitle: str, size=(1600, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, (248, 250, 252))
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, bold=True)
    subtitle_font = get_font(18)
    draw.text((60, 40), title, font=title_font, fill=(15, 23, 42))
    draw.text((60, 90), subtitle, font=subtitle_font, fill=(71, 85, 105))
    return img, draw


def save_architecture_diagram() -> Path:
    path = DIAGRAM_DIR / "system_architecture.png"
    img, draw = new_canvas(
        "MAZA End-to-End System Architecture",
        "Content is managed in the CRM, stored and exposed by Strapi, and consumed by the Expo mobile app.",
    )
    font = get_font(24, bold=True)
    small = get_font(18)
    draw_box(draw, (90, 220, 430, 380), "MAZA CRM\nVite + React + TanStack", (255, 244, 214), (217, 119, 6), font)
    draw_box(draw, (610, 180, 990, 420), "Strapi Backend\nContent types, auth, analytics,\ncustom routes, AI survey", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (1170, 220, 1510, 380), "Expo App\nReact Native + Expo Router", (220, 252, 231), (22, 163, 74), font)
    draw_box(draw, (610, 520, 990, 730), "Supporting Services\nPostgres/SQLite, SendGrid SMTP,\nSMS provider, OpenAI, Yoma,\nSentry, PostHog, file uploads", (241, 245, 249), (100, 116, 139), font)
    arrow(draw, (430, 300), (610, 300))
    arrow(draw, (990, 300), (1170, 300))
    arrow(draw, (800, 420), (800, 520))
    label(draw, (470, 255), "CRUD / upload / publish", small)
    label(draw, (1015, 255), "REST API / JWT", small)
    label(draw, (835, 460), "infra integrations", small)
    img.save(path)
    return path


def save_mobile_flow_diagram() -> Path:
    path = DIAGRAM_DIR / "mobile_user_flow.png"
    img, draw = new_canvas(
        "Expo App User Flow",
        "Primary mobile journey from onboarding through learning, opportunities, and profile management.",
    )
    font = get_font(22, bold=True)
    small = get_font(18)
    boxes = [
        ((70, 220, 280, 330), "Start"),
        ((340, 220, 610, 330), "Login / Registration"),
        ((670, 220, 960, 330), "Self-Assessment"),
        ((1020, 220, 1310, 330), "Interest Customization"),
        ((70, 500, 320, 610), "Home"),
        ((400, 500, 650, 610), "My Journey"),
        ((730, 500, 980, 610), "Opportunities"),
        ((1060, 500, 1310, 610), "Profile"),
    ]
    for coords, text in boxes:
        draw_box(draw, coords, text, (255, 255, 255), (59, 130, 246), font)
    arrow(draw, (280, 275), (340, 275))
    arrow(draw, (610, 275), (670, 275))
    arrow(draw, (960, 275), (1020, 275))
    for x in [195, 525, 855, 1185]:
        arrow(draw, (x, 330), (x, 500))
    label(draw, (130, 420), "tab navigator", small)
    img.save(path)
    return path


def save_publish_flow_diagram() -> Path:
    path = DIAGRAM_DIR / "content_publish_flow.png"
    img, draw = new_canvas(
        "Content Publishing Flow",
        "How a new course or journey moves from authoring in the CRM to delivery inside the mobile app.",
    )
    font = get_font(22, bold=True)
    small = get_font(18)
    xs = [60, 360, 660, 960, 1260]
    labels = [
        "Author content\nin CRM",
        "Upload media\n/images/files",
        "Persist draft\nin Strapi",
        "Publish / expose\nREST endpoints",
        "Consume in Expo\nqueries and screens",
    ]
    colors = [
        ((255, 244, 214), (217, 119, 6)),
        ((254, 242, 242), (220, 38, 38)),
        ((219, 234, 254), (37, 99, 235)),
        ((224, 231, 255), (79, 70, 229)),
        ((220, 252, 231), (22, 163, 74)),
    ]
    for i, x in enumerate(xs):
        fill, outline = colors[i]
        draw_box(draw, (x, 290, x + 250, 450), labels[i], fill, outline, font)
        if i < len(xs) - 1:
            arrow(draw, (x + 250, 370), (xs[i + 1], 370))
    label(draw, (735, 215), "status=draft or published", small)
    img.save(path)
    return path


def save_learning_flow_diagram() -> Path:
    path = DIAGRAM_DIR / "learning_progress_flow.png"
    img, draw = new_canvas(
        "Learning Progress and Certification Flow",
        "User-course records track enrollment, content completion, quizzes, reviews, and certificate issuance.",
    )
    font = get_font(20, bold=True)
    small = get_font(16)
    nodes = [
        ((90, 220, 350, 330), "Browse course"),
        ((450, 220, 710, 330), "Start course\nPOST /user-courses"),
        ((810, 220, 1070, 330), "Track module/content\nPUT progress endpoints"),
        ((1170, 220, 1430, 330), "Quiz / final test"),
        ((450, 520, 710, 630), "Finished state"),
        ((810, 520, 1070, 630), "Certificate generation"),
        ((1170, 520, 1430, 630), "Review / forum / analytics"),
    ]
    for coords, text in nodes:
        draw_box(draw, coords, text, (255, 255, 255), (14, 116, 144), font)
    arrow(draw, (350, 275), (450, 275))
    arrow(draw, (710, 275), (810, 275))
    arrow(draw, (1070, 275), (1170, 275))
    arrow(draw, (1300, 330), (1300, 520))
    arrow(draw, (1170, 575), (1070, 575))
    arrow(draw, (810, 575), (710, 575))
    label(draw, (905, 405), "course / module completion", small)
    img.save(path)
    return path


def save_deployment_diagram() -> Path:
    path = DIAGRAM_DIR / "deployment_topology.png"
    img, draw = new_canvas(
        "Deployment Topology",
        "Observed deployment model across web admin, backend API, and mobile distribution.",
    )
    font = get_font(22, bold=True)
    small = get_font(18)
    draw_box(draw, (80, 220, 420, 370), "Vercel\nCRM SPA\nmaza-crm", (255, 244, 214), (217, 119, 6), font)
    draw_box(draw, (80, 470, 420, 620), "Vercel\nExpo Web Export\norg.maza.app (web)", (219, 234, 254), (37, 99, 235), font)
    draw_box(draw, (630, 220, 1010, 620), "Strapi Runtime\nconfig/server.ts\nconfig/database.ts\ncustom plugin overrides", (255, 255, 255), (100, 116, 139), font)
    draw_box(draw, (1210, 220, 1520, 370), "EAS Build\nAndroid / iOS binaries", (220, 252, 231), (22, 163, 74), font)
    draw_box(draw, (1210, 470, 1520, 620), "External Services\nSMTP, SMS, OpenAI,\nYoma, Sentry, PostHog", (241, 245, 249), (71, 85, 105), font)
    arrow(draw, (420, 295), (630, 295))
    arrow(draw, (420, 545), (630, 545))
    arrow(draw, (1010, 295), (1210, 295))
    arrow(draw, (1010, 545), (1210, 545))
    label(draw, (500, 245), "admin API", small)
    label(draw, (500, 495), "mobile/web API", small)
    img.save(path)
    return path


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_table(document: Document, headers: list[str], rows: Iterable[Iterable[str]]):
    rows = list(rows)
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True)
        shade_cell(table.cell(0, i), "DCE6F1")
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value)
    return table


def add_code_block(document: Document, text: str):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(30, 64, 175)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor(51, 65, 85)

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(9.5)


def add_cover(document: Document):
    p = document.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("MAZA Platform Engineering Documentation")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 0.2\nGenerated from the current repositories on 2026-04-20")
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "This document supersedes the narrative content in "
        f"{SOURCE_DOC.name} and aligns the documentation with the current Expo app, CRM, and Strapi backend."
    )

    logo = ROOT / "assets" / "images" / "maza-logo.png"
    if logo.exists():
        document.add_picture(str(logo), width=Inches(2.3))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_numbered(
        document,
        [
            f"Expo mobile and web app: {DISPLAY_ROOT}",
            f"Strapi backend: {DISPLAY_BACKEND}",
            f"CRM/admin console: {DISPLAY_CRM}",
        ],
    )
    document.add_page_break()


def add_overview(document: Document, package_json: dict, app_json: dict):
    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "MAZA is a three-part platform composed of an Expo-based mobile application, a Strapi backend, "
        "and a CRM web application. The CRM is the operational interface used to create and manage content. "
        "Strapi is the system of record and integration layer. The Expo app is the end-user delivery channel "
        "for learning content, career guidance, opportunities, profile data, and certificates."
    )
    add_bullets(
        document,
        [
            f"Expo app version observed in app.json: {app_json['expo']['version']}",
            f"Expo SDK observed in package.json: {package_json['dependencies']['expo']}",
            f"Primary API base URL in mobile source: https://strapi.mazas.org/api",
            "CRM uses an environment-driven base URL through VITE_API_URL.",
            "Jobs are fetched from an external WordPress API at https://www.emprego.co.mz/wp-api.",
        ],
    )

    document.add_heading("2. System Context", level=1)
    document.add_paragraph(
        "At runtime, the platform behaves as a content supply chain. Administrators and content teams use the CRM "
        "to create and publish subjects, journeys, and courses. The CRM writes to Strapi over authenticated REST "
        "endpoints. The Expo app reads published data from Strapi, writes user progress and profile updates back to "
        "Strapi, and uses a separate external jobs feed for opportunities."
    )


def add_repo_responsibilities(document: Document):
    document.add_heading("3. Repository Responsibilities", level=1)
    add_table(
        document,
        ["Repository", "Primary Role", "Key Technologies", "Operational Notes"],
        [
            [
                "org.maza.app",
                "End-user mobile and web application",
                "Expo, React Native, Expo Router, React Query, Axios",
                "Consumes Strapi and external jobs API; distributed through EAS and exportable to web.",
            ],
            [
                "maza-strapi-backend",
                "Content management API, auth, analytics, and business logic",
                "Strapi 5, Node.js, custom users-permissions extension",
                "Provides custom routes for courses, analytics, forum, user progress, AI survey, OTP, and certificates.",
            ],
            [
                "maza-crm",
                "Admin and analytics interface for MAZA operators",
                "Vite, React, TanStack Router, React Query, Axios",
                "Writes directly to Strapi using JWT auth and media uploads.",
            ],
        ],
    )


def add_diagrams(document: Document, diagrams: dict[str, Path]):
    document.add_heading("4. Architecture Diagrams", level=1)
    captions = {
        "system": "Figure 1. End-to-end system architecture",
        "mobile": "Figure 2. Primary mobile user flow",
        "publish": "Figure 3. Content publishing and delivery flow",
        "learning": "Figure 4. Course progress and certification flow",
        "deployment": "Figure 5. Deployment topology",
    }
    order = ["system", "mobile", "publish", "learning", "deployment"]
    for key in order:
        document.add_picture(str(diagrams[key]), width=Inches(6.7))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(captions[key])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True


def add_mobile_architecture(document: Document, routes: list[str], package_json: dict):
    document.add_heading("5. Expo App Architecture", level=1)
    document.add_paragraph(
        "The mobile app is built as an Expo Router application with a root stack and a tab navigator. "
        "It uses React Query for server state, AsyncStorage for session persistence, Sentry for crash and performance "
        "monitoring, and PostHog for product analytics."
    )
    add_bullets(
        document,
        [
            "Root providers include QueryClientProvider, ThemeProvider, SafeAreaProvider, Sentry, and PostHog.",
            "Authentication state is cached in AsyncStorage under the @user key.",
            "The tab shell exposes Home, My Journey, Opportunities, and Profile.",
            "The app includes dedicated onboarding, self-assessment, search, room, jobs, challenges, missions, and user profile flows.",
        ],
    )

    document.add_heading("5.1 Key Mobile Modules", level=2)
    add_table(
        document,
        ["Module", "Representative Files", "Responsibility"],
        [
            ["Navigation", "app/_layout.tsx, app/(tabs)/_layout.tsx", "Bootstraps providers and controls root and tab navigation."],
            ["Authentication", "app/services/authService.ts, app/hooks/useAuthMutations.ts", "OTP login, email login, registration, password reset, profile session management."],
            ["Catalog and Learning", "services/catalog.ts, app/room/*", "Course detail, enrollment, progress tracking, quizzes, reviews, forum, certificates."],
            ["Home and Discovery", "services/home.ts, app/(tabs)/index.tsx, app/search/*", "Home feed, new/popular/suggested courses, subject discovery."],
            ["Opportunities", "services/jobsApi.ts, app/(tabs)/opportunities.tsx, app/jobs/*", "Job discovery flow backed by external emprego.co.mz API."],
            ["Onboarding", "services/self-assessment.ts, services/survey.ts, util/onboarding.ts", "AI-guided self-assessment, route gating, interest customization."],
            ["Observability", "utils/analytics.ts, @sentry/react-native", "User identification, screen tracking, runtime telemetry."],
        ],
    )

    document.add_heading("5.2 Observed Route Inventory", level=2)
    route_rows = [[route, describe_route(route)] for route in routes]
    add_table(document, ["Route File", "Purpose"], route_rows[:26])
    if len(route_rows) > 26:
        document.add_paragraph(
            f"The complete route inventory contains {len(route_rows)} files and is reproduced in Annex A."
        )

    document.add_heading("5.3 Mobile Dependencies", level=2)
    dep_rows = []
    for name in [
        "expo",
        "expo-router",
        "@tanstack/react-query",
        "axios",
        "@sentry/react-native",
        "posthog-react-native",
        "react-native-pdf",
        "react-native-youtube-iframe",
    ]:
        dep_rows.append([name, package_json["dependencies"][name]])
    add_table(document, ["Dependency", "Observed Version"], dep_rows)


def describe_route(route: str) -> str:
    descriptions = {
        "_layout.tsx": "Application root layout and provider tree.",
        "modal.tsx": "Generic modal screen.",
        "(tabs)/index.tsx": "Home tab.",
        "(tabs)/courses.tsx": "Learning journey tab.",
        "(tabs)/opportunities.tsx": "Jobs/opportunities tab.",
        "(tabs)/profile.tsx": "Profile tab.",
        "login/index.tsx": "Phone-first login entry.",
        "login/login-email.tsx": "Email/password login.",
        "login/otp.tsx": "OTP verification.",
        "onboarding/self-assessment.tsx": "AI self-assessment interview.",
        "onboarding/survey.tsx": "Survey-driven onboarding.",
        "start/customize.tsx": "Interest personalization.",
        "journeys/index.tsx": "Journey listing.",
        "journeys/[id].tsx": "Journey detail.",
        "jobs/search.tsx": "Job search results.",
        "jobs/[slug].tsx": "Job detail.",
        "room/lessons.tsx": "Course lesson navigation.",
        "room/watch.tsx": "Video/content playback.",
        "room/quiz.tsx": "Quiz and final test handling.",
        "room/certificate.tsx": "Certificate view/download.",
    }
    return descriptions.get(route, "Feature screen or supporting route.")


def add_integration_section(document: Document):
    document.add_heading("6. Integration Architecture", level=1)
    document.add_paragraph(
        "The integration model is API-first. Both the CRM and Expo app call Strapi over REST. "
        "The mobile app combines authenticated and unauthenticated endpoints, while the CRM authenticates operators "
        "and then performs content CRUD, publishing, upload, analytics, and user management operations."
    )

    document.add_heading("6.1 Mobile to Strapi", level=2)
    add_table(
        document,
        ["Capability", "Endpoint Pattern", "Notes"],
        [
            ["Profile retrieval", "GET /users/me", "JWT-protected; used by onboarding, profile, and session refresh."],
            ["Email/OTP login", "POST /auth/login", "Supports password and OTP variants in the same route."],
            ["OTP request", "POST /otps", "Phone-first authentication bootstrap."],
            ["Registration", "POST /users", "Creates mobile user and may trigger Yoma account creation asynchronously."],
            ["Survey/self-assessment", "POST /users-permissions/survey", "Custom AI-backed conversational assessment route."],
            ["Interests", "POST/DELETE /users-permissions/interests", "Connects and disconnects user interests."],
            ["Catalog", "GET /courses, /courses/{id}, /subjects, /journeys", "Drives home, search, and journey flows."],
            ["Enrollment/progress", "POST/GET/PUT /user-courses...", "Tracks course state, content completion, and quiz grades."],
            ["Community", "GET/POST /courses/{id}/forum...", "Forum comments and replies."],
            ["Reviews", "GET/POST /reviews", "Course reviews."],
            ["Certificates", "GET /certificates", "Certificate retrieval after course completion."],
        ],
    )

    document.add_heading("6.2 CRM to Strapi", level=2)
    add_table(
        document,
        ["CRM Area", "Endpoint Pattern", "Purpose"],
        [
            ["Authentication", "POST /auth/crm", "Role-restricted login for crm_admin, crm_editor, crm_analytics."],
            ["Courses", "GET /courses/all, GET /courses/{id}, POST /courses, PUT /courses/{id}", "Create/edit/publish learning content."],
            ["Media", "POST /upload", "Uploads course images and generic assets."],
            ["Journeys", "GET/POST/PUT /journeys", "Learning pathway management."],
            ["Subjects", "GET/POST/PUT /subjects", "Category and subject management."],
            ["Analytics", "GET /events/stats/*, /courses/stats/*", "Dashboard counts and trend data."],
            ["Users", "GET/PUT/DELETE /users...", "Profile inspection, role changes, and account blocking/deactivation."],
        ],
    )

    document.add_heading("6.3 External Services", level=2)
    add_table(
        document,
        ["Service", "Repository Touchpoint", "Role"],
        [
            ["emprego.co.mz WordPress API", "Expo app services/jobsApi.ts", "External job opportunities feed."],
            ["OpenAI API", "Strapi users-permissions extension", "Conversational self-assessment and recommendation engine."],
            ["Yoma API", "Strapi users-permissions extension", "External youth account creation / identity linkage."],
            ["SMS Provider (MozeSMS / Wei / Twilio)", "Strapi OTP and environment config", "OTP delivery and phone verification workflows."],
            ["SendGrid / SMTP", "Strapi plugins.ts", "Email delivery and system messaging."],
            ["PostHog", "Expo app", "Product analytics and user/session identification."],
            ["Sentry", "Expo app", "Crash, trace, and performance monitoring."],
        ],
    )


def add_backend_section(document: Document):
    document.add_heading("7. Strapi Backend Architecture", level=1)
    document.add_paragraph(
        "The backend is a Strapi 5 application extended through custom APIs, content-type lifecycle hooks, "
        "and a substantial override of the users-permissions plugin. It functions as both a CMS and the core "
        "application backend for MAZA."
    )
    add_bullets(
        document,
        [
            "Primary custom APIs include analytics, certificate, course, event, journey, OTP, review, subject, and user-course.",
            "The users-permissions plugin is extended for mobile login, CRM login, AI self-assessment, interests management, and admin-controlled user updates/deactivation.",
            "The backend logs key product events such as account creation, login success/failure, profile updates, and assessment completion.",
            "Database configuration supports sqlite and postgres; local development defaults to sqlite while production-oriented settings include PostgreSQL parameters.",
        ],
    )

    document.add_heading("7.1 Core Content and Domain Objects", level=2)
    add_table(
        document,
        ["Domain Object", "Purpose", "Observed Related Files"],
        [
            ["Course", "Primary learning unit with content modules, forum, reviews, and statistics.", "src/api/course/*"],
            ["Subject", "Topic/category layer used for organization and filtering.", "src/api/subject/*"],
            ["Journey", "Learning pathway tying together categories/courses.", "src/api/journey/*"],
            ["User-Course", "Enrollment, progress, quiz completion, and state tracking.", "src/api/user-course/*"],
            ["Certificate", "Issued after completion and exposed to the mobile app.", "src/api/certificate/*"],
            ["Event", "Analytics and audit event storage.", "src/api/event/*"],
            ["OTP", "Authentication helper for phone-based access.", "src/api/otp/*"],
        ],
    )

    document.add_heading("7.2 Backend Runtime Configuration", level=2)
    add_bullets(
        document,
        [
            "config/server.ts exposes host, port, URL, proxy, and Strapi app keys.",
            "config/plugins.ts configures users-permissions JWT lifetime to 180 days and email delivery via nodemailer/SMTP.",
            "config/middlewares.ts loosens frameguard and adds custom CSP allowances for localhost and S3 media.",
            "The AWS S3 upload provider is present but currently commented in config/plugins.ts, which suggests either local or alternate upload behavior in the observed state.",
        ],
    )

    document.add_heading("7.3 AI Self-Assessment Implementation", level=2)
    document.add_paragraph(
        "The self-assessment is handled server-side inside the users-permissions extension. The server builds a "
        "developer prompt in Portuguese, injects demographic context from the authenticated user, and calls the "
        "OpenAI Chat Completions API. When the assistant emits the sentinel token ###SURVEY_DONE###, Strapi stores "
        "the final recommendation back onto the user profile and logs an ASSESS_SURVEY event."
    )


def add_crm_section(document: Document):
    document.add_heading("8. CRM Architecture", level=1)
    document.add_paragraph(
        "The CRM is a Vite-based React SPA organized around TanStack Router and React Query. It is the operational "
        "surface for content creation, publication, uploads, analytics dashboards, and user administration. "
        "Its API client is fully environment-driven through VITE_API_URL and injects the CRM JWT token on each request."
    )
    add_table(
        document,
        ["Area", "Representative Files", "Responsibility"],
        [
            ["API Client", "src/lib/api-client.ts", "Base URL, auth token injection, and 401 handling."],
            ["Course Management", "src/http/courses/*, routes/dashboard/courses/*", "Course CRUD, publish/unpublish, image uploads, course builder UX."],
            ["Journey Management", "src/http/journeys/*", "Journey list/detail and mutations."],
            ["Subject Management", "src/http/subjects/*", "Subject CRUD and subject analytics."],
            ["Analytics", "src/http/analytics/queries.ts", "Signups, active users, course engagement, password resets, certificate downloads, trends."],
            ["User Management", "src/http/users/*", "User lookup, update, role changes, and deactivation workflows."],
        ],
    )
    document.add_paragraph(
        "The CRM README still contains template TanStack content, so the implementation source is currently a more reliable "
        "source of truth than the repository README."
    )


def add_deployment_section(document: Document):
    document.add_heading("9. Deployment and Environments", level=1)
    document.add_heading("9.1 Expo App", level=2)
    add_bullets(
        document,
        [
            "app.json defines bundle/package identifier org.maza.app for both iOS and Android.",
            "EAS profiles include development, staging, preview, and production.",
            "vercel.json exports a static web build using `npx expo export --platform web` with SPA rewrites.",
            "Observed production API endpoint in source is hardcoded to https://strapi.mazas.org/api.",
        ],
    )

    document.add_heading("9.2 CRM", level=2)
    add_bullets(
        document,
        [
            "The CRM is configured as a Vercel-style SPA with a catch-all rewrite to /index.html.",
            "Environment variable VITE_API_URL points the CRM to the Strapi API base.",
            "Production build command is `pnpm build` and local development runs with `pnpm dev` on port 3000.",
        ],
    )

    document.add_heading("9.3 Strapi Backend", level=2)
    add_bullets(
        document,
        [
            "The backend exposes develop, start, build, and deploy scripts plus a local setup script.",
            "config/database.ts supports sqlite for local work and postgres for production deployments.",
            "Deployment requires APP_KEYS, JWT secrets, database connection values, SMTP/SMS config, and API keys for OpenAI and Yoma if those integrations are enabled.",
        ],
    )

    document.add_heading("9.4 Minimal Release Workflow", level=2)
    add_numbered(
        document,
        [
            "Update or create content in the CRM.",
            "Publish content through Strapi-backed endpoints so the mobile app can see the published variant.",
            "Verify catalog and journey payloads through Strapi or the CRM list views.",
            "Build mobile binaries with EAS for Android/iOS as needed.",
            "Deploy CRM and backend changes to their hosting targets.",
            "Smoke-test mobile login, onboarding, course enrollment, opportunities, and certificate retrieval against production APIs.",
        ],
    )


def add_ops_section(document: Document):
    document.add_heading("10. Operations, Monitoring, and Security", level=1)
    add_table(
        document,
        ["Concern", "Observed Mechanism", "Notes"],
        [
            ["Crash monitoring", "Sentry in Expo root layout", "Initialized outside development mode."],
            ["Product analytics", "PostHog in Expo app", "User identified by documentId plus name/identifier metadata."],
            ["JWT auth", "Strapi users-permissions", "CRM and mobile flows use different login entry points."],
            ["Session storage", "AsyncStorage in Expo, local auth storage in CRM", "Mobile stores user session under @user."],
            ["Rate limiting", "users-permissions rateLimit middleware", "Applied to auth routes configured in the extension."],
            ["Email and SMS", "SMTP and SMS providers in Strapi config", "Requires production-grade credentials and sender identities."],
        ],
    )


def add_risks_section(document: Document):
    document.add_heading("11. Observed Risks and Alignment Notes", level=1)
    add_bullets(
        document,
        [
            "The mobile README is outdated. It lists older versions and historical API endpoints that do not match current source files.",
            "The mobile API layer hardcodes the production Strapi URL in services/api.ts while .env.example advertises EXPO_PUBLIC_BASE_URL. This creates configuration drift risk.",
            "util/onboarding.ts checks whether `survey` is an array, but the backend self-assessment persists survey recommendations as a string. This may cause onboarding completion checks to misbehave.",
            "The CRM README is still template boilerplate and does not document the actual domain model or deployment process.",
            "Course publishing in CRM currently sets a hardcoded publishedAt timestamp in one mutation path, which should be replaced with a runtime timestamp or dedicated publish route usage.",
            "Several files show encoding artifacts in Portuguese strings. Documentation, UI validation, and API messages would benefit from a repository-wide encoding cleanup.",
        ],
    )


def add_mermaid_appendix(document: Document, diagrams: dict[str, Path]):
    document.add_heading("12. Mermaid Appendix", level=1)
    document.add_paragraph(
        "The following Mermaid blocks can be reused in Markdown documentation, wiki pages, PRs, or engineering knowledge bases."
    )
    blocks = [
        (
            "System Architecture",
            "system",
            """flowchart LR
    CRM["MAZA CRM"]
    STRAPI["Strapi Backend"]
    APP["Expo App"]
    EXT["OpenAI / Yoma / SMTP / SMS / Analytics"]
    CRM -->|CRUD / publish / upload| STRAPI
    APP -->|JWT-authenticated reads and writes| STRAPI
    STRAPI --> EXT
    APP -->|Jobs feed| JOBS["emprego.co.mz WP API"]""",
        ),
        (
            "User Journey",
            "mobile",
            """flowchart LR
    Start --> Login
    Login --> Assessment
    Assessment --> Interests
    Interests --> Home
    Home --> Journey
    Home --> Opportunities
    Home --> Profile
    Journey --> Room
    Room --> Certificate""",
        ),
        (
            "Content Publishing",
            "publish",
            """flowchart LR
    Author["Author in CRM"] --> Upload["Upload media"]
    Upload --> Draft["Save draft in Strapi"]
    Draft --> Publish["Publish content"]
    Publish --> Query["Expo queries data"]
    Query --> Render["Render course/journey in app"]""",
        ),
        (
            "Learning Progress",
            "learning",
            """flowchart LR
    Browse --> Enroll["POST /user-courses"]
    Enroll --> Watch["Complete content"]
    Watch --> Quiz["Module quiz / final test"]
    Quiz --> Finish["Finished state"]
    Finish --> Certificate["Generate / fetch certificate"]
    Finish --> Review["Review and analytics events"]""",
        ),
    ]
    for title, diagram_key, block in blocks:
        document.add_heading(title, level=2)
        document.add_picture(str(diagrams[diagram_key]), width=Inches(6.3))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = document.add_paragraph(f"Rendered flowchart for {title}.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        add_code_block(document, block)


def add_annexes(document: Document, routes: list[str], app_json: dict, eas_json: dict):
    document.add_heading("Annex A. Full Route Inventory", level=1)
    add_table(document, ["Route File", "Purpose"], [[route, describe_route(route)] for route in routes])

    document.add_heading("Annex B. Environment Variables", level=1)
    add_table(
        document,
        ["Repository", "Variable", "Purpose"],
        [
            ["Expo app", "EXPO_PUBLIC_BASE_URL", "Declared example env var for API base URL; not currently used by services/api.ts."],
            ["Expo app", "EXPO_PUBLIC_POSTHOG_API_KEY", "Initializes PostHog analytics client."],
            ["CRM", "VITE_API_URL", "Points the CRM to Strapi API base URL."],
            ["Strapi", "SERVER_ADDRESS", "Public server URL used by config/server.ts."],
            ["Strapi", "DATABASE_* / DATABASE_URL", "Database connection configuration."],
            ["Strapi", "OPENAI_API_*", "AI survey provider configuration."],
            ["Strapi", "SMTP_* / SENDGRID_API_KEY", "Email delivery settings."],
            ["Strapi", "YOMA_*", "External youth account provisioning."],
            ["Strapi", "MOZESMS_* / TWILIO_* / WEI_*", "OTP and SMS integration settings."],
        ],
    )

    document.add_heading("Annex C. Build Profiles", level=1)
    rows = []
    for profile, config in eas_json["build"].items():
        rows.append(
            [
                profile,
                "yes" if config.get("distribution") == "internal" else "no",
                json.dumps(config.get("env", {}), ensure_ascii=False),
                json.dumps({k: v for k, v in config.items() if k in {"android", "ios", "developmentClient", "autoIncrement"}}, ensure_ascii=False),
            ]
        )
    add_table(document, ["Profile", "Internal Distribution", "Env Overrides", "Platform Notes"], rows)

    document.add_heading("Annex D. Document Metadata", level=1)
    add_bullets(
        document,
        [
            f"Source document reviewed: {DISPLAY_SOURCE_DOC}",
            f"Generated document path: {DISPLAY_OUT_DOC}",
            f"Expo app path: {DISPLAY_ROOT}",
            f"CRM path: {DISPLAY_CRM}",
            f"Backend path: {DISPLAY_BACKEND}",
            f"Observed Expo version: {app_json['expo']['version']}",
        ],
    )


def build_document():
    ensure_dirs()
    package_json = load_json(ROOT / "package.json")
    app_json = load_json(ROOT / "app.json")
    eas_json = load_json(ROOT / "eas.json")
    routes = list_routes()

    diagrams = {
        "system": save_architecture_diagram(),
        "mobile": save_mobile_flow_diagram(),
        "publish": save_publish_flow_diagram(),
        "learning": save_learning_flow_diagram(),
        "deployment": save_deployment_diagram(),
    }

    document = Document()
    configure_styles(document)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        add_page_number(section)

    add_cover(document)
    add_overview(document, package_json, app_json)
    add_repo_responsibilities(document)
    add_diagrams(document, diagrams)
    add_mobile_architecture(document, routes, package_json)
    add_integration_section(document)
    add_backend_section(document)
    add_crm_section(document)
    add_deployment_section(document)
    add_ops_section(document)
    add_risks_section(document)
    add_mermaid_appendix(document, diagrams)
    add_annexes(document, routes, app_json, eas_json)

    document.save(str(OUT_DOC))


if __name__ == "__main__":
    build_document()
